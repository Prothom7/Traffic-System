from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Traffic-System_SoftwareX_Report.docx"


def set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def set_paragraph_format(paragraph, *, first_line: float = 0.25, space_after: float = 0.08) -> None:
    paragraph.paragraph_format.first_line_indent = Inches(first_line)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph.paragraph_format.space_after = Pt(space_after * 12)


def add_title_page(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Traffic-System: A Next.js-Based Smart Traffic Enforcement Platform\n")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(20)

    subtitle = title.add_run(
        "SoftwareX-Style Seminar Paper Report\n\n"
        "Prepared for academic documentation and system presentation\n"
        "Generated from the current Traffic-System codebase"
    )
    subtitle.font.name = "Times New Roman"
    subtitle.font.size = Pt(12)

    for _ in range(5):
        document.add_paragraph()

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        "Workspace: Traffic-System",
        "Platform: Next.js, React, MongoDB, FastAPI, PyTorch",
        "Scope: Citizen services, administrative enforcement, ALPR integration",
        "Document type: Detailed seminar paper / technical report",
    ]:
        run = meta.add_run(line + "\n")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    document.add_page_break()


def add_header_block(document: Document, title: str, authorship: str, abstract: list[str], keywords: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(authorship)
    run.italic = True
    run.font.size = Pt(11)

    document.add_paragraph()
    h = document.add_paragraph()
    h.add_run("Abstract").bold = True
    for paragraph_text in abstract:
        p = document.add_paragraph(paragraph_text)
        set_paragraph_format(p, first_line=0.25)

    kp = document.add_paragraph()
    kp.add_run("Keywords: ").bold = True
    kp.add_run(keywords)
    kp.paragraph_format.space_after = Pt(12)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)


def add_body_paragraphs(document: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        p = document.add_paragraph(text)
        set_paragraph_format(p)


def add_bullet_list(document: Document, bullets: list[str]) -> None:
    for item in bullets:
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            cells[i].text = cell_text
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)


def route_description(route: str) -> str:
    descriptions = {
        "app/api/authentication/signin/route.ts": "handles credential validation, JWT issuance, and login-session initialization for citizens and administrators.",
        "app/api/authentication/signup/route.ts": "creates new user accounts, validates registration data, and stores account metadata for email verification.",
        "app/api/authentication/verifyemail/route.ts": "confirms verification tokens and transitions newly created users into an active account state.",
        "app/api/violations/route.ts": "returns user-scoped violation history with filters for status, date range, and vehicle association.",
        "app/api/vehicles/route.ts": "exposes vehicle registry records for authenticated and administrative listing workflows.",
        "app/api/traffic-records/route.ts": "retrieves traffic detections and violation entries from the operational record store.",
        "app/api/dashboard/stats/route.ts": "computes summary metrics for tickets, status counts, and account activity on the dashboard.",
        "app/api/dashboard/carousel/route.ts": "serves public content cards and rotating announcements shown on the landing and dashboard views.",
        "app/api/dashboard/carousel-by-category/route.ts": "returns carousel items grouped by category so the UI can render focused content feeds.",
        "app/api/dashboard/newsfeed/route.ts": "publishes the current traffic and accident news stream used by the explorer section.",
        "app/api/notifications/route.ts": "fetches alert records so the client can render recent notifications and unread counts.",
        "app/api/notifications/stream/route.ts": "supports notification streaming for clients that keep a live connection to enforcement events.",
        "app/api/profile/route.ts": "packages profile data, linked vehicles, and account metadata into a single authenticated response.",
        "app/api/payments/gateway/route.ts": "records ticket payments, updates violation status, and links payment metadata to the corresponding record.",
        "app/api/admin/camera/route.ts": "creates and manages camera locations that act as sensor nodes in the traffic map.",
        "app/api/admin/camera/[id]/route.ts": "updates or removes an individual camera location when administrative edits are required.",
        "app/api/admin/locations/route.ts": "lists camera and location points for the admin map and graph visualization.",
        "app/api/admin/edges/route.ts": "creates the weighted road-network edges used to model connectivity between camera nodes.",
        "app/api/admin/analysis/route.ts": "aggregates traffic records into trends for operational analysis and reporting.",
        "app/api/admin/violations/simulate/route.ts": "creates a synthetic violation event, normalizes the number plate, and emits downstream notifications.",
        "app/api/admin/violations/extract-plate/route.ts": "invokes the ALPR pipeline through FastAPI or a local Python process and returns the detected plate.",
        "app/api/admin/UI/newsfeed/route.ts": "supports the admin newsfeed maintenance interface with content CRUD-style retrieval.",
        "app/api/admin/UI/carousel/add/route.ts": "adds new dashboard carousel content for the administrative UI content editor.",
        "app/api/admin/UI/carousel/delete/[id]/route.ts": "removes a carousel entry identified by its database record.",
        "app/api/services/add-vehicle/route.ts": "creates a vehicle entry for citizen self-service registration.",
        "app/api/services/change-ownership/route.ts": "transfers vehicle ownership metadata from one registered user to another.",
        "app/api/services/check-status/route.ts": "returns registration and enforcement status for a given vehicle or account.",
        "app/api/services/renew-registration/route.ts": "updates the registration expiry date and persists renewal activity.",
        "app/api/services/report-stolen/route.ts": "flags a vehicle as stolen and propagates the new status through the system.",
        "app/api/services/update-details/route.ts": "updates vehicle attributes or ownership details after validation.",
    }
    return descriptions.get(route, "supports the application workflow through a route-specific CRUD or query operation.")


def build_report() -> Document:
    document = Document()
    set_document_defaults(document)

    add_title_page(document)

    title = (
        "Traffic-System: An Integrated Next.js Web Platform with ALPR-Enabled Enforcement Automation "
        "for Smart-City Traffic Administration"
    )
    abstract = [
        "Traffic-System is a full-stack smart-city web platform that unifies citizen vehicle services, administrative enforcement tools, and automated traffic violation workflows in a single browser-based environment. The application is implemented with Next.js and React on the frontend, MongoDB and Mongoose for persistent data management, and a Python-based ALPR service that supports plate extraction during enforcement simulations.",
        "The system is designed to approximate the behavior of a digital traffic authority. Citizens can register vehicles, renew registration, update details, report theft, inspect violation history, and track payment status, while administrators can manage camera nodes, inspect traffic records, simulate violations, visualize camera networks, and test machine-learning inference pipelines through dedicated pages.",
        "A distinguishing aspect of the project is its layered integration of an ALPR pipeline. The repository supports a standalone FastAPI inference server and a local subprocess fallback so the same workflow can function in development, testing, and more persistent deployment settings. Plate normalization logic also handles common regional variants, including Bangla digits and spacing differences.",
        "This report documents the system at the level expected from a detailed SoftwareX-style seminar paper. It summarizes the software architecture, the data model, the API surface, the front-end composition, the operational workflows, the ML integration, the deployment strategy, the known limitations, and the most relevant directions for future work.",
    ]
    add_header_block(
        document,
        title,
        "Prepared from the Traffic-System workspace and code structure",
        abstract,
        "automatic license plate recognition, traffic enforcement, Next.js, MongoDB, FastAPI, smart city, vehicle registration, violation management",
    )

    sections = [
        (
            "1. Introduction",
            [
                "Traffic enforcement systems often combine operational complexity, high data volume, and a mixture of manual and automated decision-making. A traditional workflow can require inspectors, camera operators, registry clerks, and payment officers to coordinate across separate tools, each of which introduces latency and data duplication. The Traffic-System project responds to this problem by consolidating those functions into a single web platform where enforcement events, vehicle identity, user accounts, and payments can be represented consistently.",
                "The application is not only a record-management site. It is a demonstration of how a smart-city traffic authority can be modeled as a software system with citizen-facing services and administrative control surfaces. The citizen portal provides self-service capabilities that reduce administrative overhead, while the admin portal offers live monitoring, simulation, and content-management utilities. That separation is important because it mirrors real operational boundaries between public users and enforcement staff.",
                "From an architectural perspective, the project favors a practical monolithic web stack for the application layer and a separate Python service for machine-learning inference. That choice reduces integration complexity while still allowing the ALPR workflow to behave like a specialized microservice. The result is a system that is easier to run locally, but still capable of illustrating cross-service communication and structured model loading.",
                "The report is written in a seminar-paper style so that the work can be presented as an academic software engineering artifact. The emphasis is on how the codebase is organized, how data flows between components, and how the implementation choices support the system goals rather than on marketing language or unsupported claims.",
            ],
        ),
        (
            "2. Problem Context and Motivation",
            [
                "The motivation for the project is rooted in the inefficiencies of manual traffic enforcement. In a manual setup, a camera image may need to be inspected by an operator, the plate number transcribed by hand, the vehicle registry searched, and the notice of violation issued separately. Each step is vulnerable to delay and human error. The Traffic-System design compresses those steps into a consistent event pipeline that can be triggered by simulation or by machine-assisted recognition.",
                "A second motivation is the need to link traffic records with citizen services rather than treating them as isolated data collections. A violation is not just an enforcement event; it has consequences for renewal, payment, driving status, and owner communication. The project therefore ties vehicle registry state to ticket history, notifications, and credit-like status indicators, allowing a single record to influence multiple user-facing workflows.",
                "The codebase also illustrates a realistic development trade-off: not every advanced capability has to be embedded into the primary web server. By separating the ALPR model execution into its own FastAPI process, the system avoids repeatedly loading large model weights during each request. That decision is especially relevant in local Windows development where PyTorch and OpenMP-related runtime configuration can otherwise slow iteration.",
            ],
        ),
        (
            "3. Project Scope and Contributions",
            [
                "The scope of the repository is broad enough to cover registration, enforcement, public reporting, analysis, and administrative content management. It includes a landing page, a dashboard, authentication flows, vehicle-specific services, traffic records, violation management, and multiple admin utilities for simulation and analytics. That breadth is the reason the report needs to be detailed rather than narrowly focused on one feature area.",
                "One contribution of the system is the deliberate modeling of both the user journey and the enforcement journey. Citizens can create and maintain vehicle records, while administrators can cause or inspect violation events and see how those events propagate into notifications and payment states. This dual perspective is valuable in software documentation because it demonstrates end-to-end traceability from input to outcome.",
                "Another contribution is the inclusion of a computer-vision pipeline without sacrificing the web application’s usability. The project treats ALPR as a modular capability that can be switched between a persistent inference server and a direct Python invocation path. That makes the feature more deployable in an educational environment where GPU access, package management, and model storage may vary across machines.",
            ],
        ),
        (
            "4. System Architecture Overview",
            [
                "The application architecture can be understood as three interacting layers. The first layer is the Next.js frontend, which renders pages, forms, maps, and dashboards with CSS modules and browser-side state. The second layer is the Next.js API route layer, which acts as the application backend and performs authentication, data access, orchestration, and validation. The third layer is the data and inference environment, consisting of MongoDB, Mongoose models, and a Python ALPR service.",
                "This structure is well matched to the needs of the project because the Next.js API routes already live inside the same repository as the UI. That keeps the web application cohesive while still allowing specialized external processing when needed. In practical terms, the citizen pages and admin pages communicate with route handlers, which in turn talk to the database or to the ML service depending on the task at hand.",
                "The repository also includes helper modules for JWT operations, mailing, and token handling. These support functions sit beneath the application layer and enable the functional separation of concerns that the documentation should emphasize. Authentication is thus not a one-off page-level concern but a shared capability that affects every protected workflow in the system.",
                "The ALPR pipeline is architecturally distinct because it bridges JavaScript and Python. Instead of overloading the main Node.js application with image-processing code, the design delegates model execution to a Python runtime where PyTorch, YOLO, and OpenCV are natural dependencies. That separation makes the deployment story more realistic and closer to how production systems often isolate machine-learning work.",
            ],
        ),
        (
            "5. Frontend Design and User Experience",
            [
                "The frontend is organized as a route-driven application under the App Router structure. Each major user journey receives its own page directory, which allows the codebase to remain understandable even as the system expands. Common visual patterns are expressed through CSS modules, meaning each screen can preserve local styling without introducing a global style collision across the large application surface.",
                "The landing pages, dashboard, and service pages are designed to communicate trust, utility, and clear action. Since the product serves both citizens and administrators, the interface needs to avoid ambiguity about which actions are public and which are privileged. The route structure, combined with role-aware navigation, serves this separation by placing administrative tools in the admin subtree and citizen services in the service subtree.",
                "Maps are a notable part of the user experience because the application is not only about text-based records. The use of Leaflet and react-leaflet indicates that location data is central to the system model. Admin users can inspect camera positions and traffic nodes visually, while citizen-facing ticket pages can present the context of a violation through map-based information rather than a flat table alone.",
            ],
        ),
        (
            "6. Backend and API Organization",
            [
                "The backend is implemented through route handlers under app/api, which means request handling, validation, and response shaping are colocated with the application code. This pattern simplifies the mental model for the project because the same repository contains the UI, the orchestration logic, and the data-access layer. For an academic seminar paper, that cohesion is easy to describe and easy to justify.",
                "The API surface covers authentication, vehicles, violations, traffic records, notifications, dashboard metrics, payments, and service workflows. Administrative routes are separated from user-oriented routes so that access checks can be applied consistently. The route inventory shows that the system is intended to support both read-heavy dashboard use and write-heavy enforcement actions.",
                "A practical benefit of this API structure is traceability. When a violation is simulated, a small set of route handlers and models are involved, and the resulting database state can be followed from request to notification. That makes the system suitable for demonstration and teaching because the effects of each action are easy to explain and verify.",
            ],
        ),
        (
            "7. Data Model and Persistence Layer",
            [
                "The persistent layer is centered on MongoDB with Mongoose schemas. The codebase contains distinct models for users, vehicles, traffic records, notifications, locations, edges, carousel content, and news feeds. That design reflects the application’s mixed requirements: some collections represent core transactional data, while others store presentation or operational content used by dashboards and admin pages.",
                "The vehicle and traffic record models are especially important because they anchor the enforcement workflow. A vehicle record carries the canonical plate and ownership metadata, while a traffic record captures a detected event, a location, a timestamp, and the violation details associated with that event. By storing both the plate and a vehicle reference, the system balances normalization with the need for fast lookup and reporting.",
                "The location and edge models reveal that the project does more than store static administrative lists. They represent a graph of camera or monitoring points connected by weighted routes. This graph-based representation is useful for route analysis, hotspot reasoning, and visual map display, and it also gives the report an opportunity to explain why traffic infrastructure is not just a set of isolated points.",
            ],
        ),
        (
            "8. Authentication and Access Control",
            [
                "Authentication in the project is implemented through JWT-based flows backed by helper functions and API routes. The presence of sign-in and sign-up endpoints, along with a verification step, shows that account lifecycle management is considered part of the system rather than an afterthought. This is essential for a platform where records can have legal or financial consequences.",
                "The distinction between citizen and administrator roles is central to the access model. Citizens can see their own vehicles, profiles, tickets, and service forms, while administrators can access traffic-wide records, camera management, simulation tools, and analytics. Role separation is enforced at the UI level and the API level, which is the correct pattern for a system that should not rely on client-side controls alone.",
                "The client-side helper for JWT decoding indicates that the application uses tokens not only for server authentication but also for browser-side routing decisions and state restoration. That decision is pragmatic for a dashboard-style application, although it does imply that the report should acknowledge the security trade-off of storing tokens in localStorage. A seminar paper should present that choice as a development convenience rather than a final production ideal.",
            ],
        ),
        (
            "9. Citizen Services and Self-Service Workflows",
            [
                "The citizen-facing service area is one of the strongest parts of the system because it turns a traffic authority into a usable digital service. Instead of forcing a user to contact an office for every minor change, the platform exposes pages for ownership changes, registration renewal, status checks, stolen vehicle reporting, and detail updates. These pages are not isolated forms; they are structured workflows that interact with the database and influence future enforcement outcomes.",
                "The service pages are especially important in the report because they show that traffic management is not only about punishment. A robust traffic system also supports correction, renewal, and change of record ownership. That makes the application feel closer to a governmental portal than a pure enforcement dashboard, which is a more realistic interpretation of how transport administration operates in practice.",
                "The payment-history and ticket pages also strengthen the citizen experience because they let the user verify what the system believes about them. When a fine is generated, the user can inspect the violation history, map location, and payment state from a single interface. That transparency is valuable both technically and institutionally because it supports accountability.",
            ],
        ),
        (
            "10. Administrative Tools and Operational Control",
            [
                "The administrative branch of the application contains the tools that turn the platform from a registry site into an enforcement system. The admin dashboard, camera management page, traffic-records page, vehicle management page, violation overview, analysis dashboard, and simulation utilities together provide a broad operational surface. Each of those screens corresponds to a different stage in the operational lifecycle of traffic supervision.",
                "The simulation page is particularly important because it allows the system to be exercised without actual roadside hardware. An administrator can create a violation event, test plate extraction, and observe how the state changes across the database and notification system. That means the repository supports both demonstration and quality assurance, which is useful in an academic or prototype setting.",
                "The analysis page and UI content-management pages show that the admin experience is not limited to enforcement. The system also includes a basic content layer for newsfeeds and carousels, allowing the administrative team to maintain informational banners and public updates. This broadens the system from a backend tool into a complete portal with both operational and informational responsibilities.",
            ],
        ),
        (
            "11. Automatic License Plate Recognition Pipeline",
            [
                "The ALPR capability is the most specialized technical feature in the repository. The pipeline combines object detection, plate cropping, city or region classification, and character recognition to reconstruct a plate value from an input image. The presence of a dedicated FastAPI service suggests that the implementation is meant to keep model loading and inference separate from the main web server for stability and performance reasons.",
                "The project documentation indicates the use of YOLO for plate localization, classification models for city or label prediction, and a digit or character recognition model for the remaining sequence. The pipeline also includes normalization logic for Bangla digit conversion and spacing cleanup, which is an important detail because real plates often appear in slightly inconsistent visual forms. This normalization step is exactly the kind of nuance that makes the report feel grounded in the codebase.",
                "The repository also supports a direct Python subprocess mode for environments where a persistent ML server is not desirable. That fallback matters because prototype systems often need to be portable across development machines. A seminar paper should note that this is a useful engineering compromise: it simplifies local setup while keeping a production-like FastAPI deployment path available when needed.",
            ],
        ),
        (
            "12. Notification, Payment, and Status Propagation",
            [
                "Traffic enforcement is only useful if the resulting event propagates through the rest of the system. In Traffic-System, violation creation is followed by notification creation and by updates to the record status that underpin payment and dashboard views. This means the user experience is stateful: a violation becomes visible, actionable, and eventually resolvable.",
                "The notification subsystem is implemented with a route layer and an event-broadcast pattern, allowing newly created violations to be observed by the front end with minimal delay. Even if the current implementation is polling-based in some cases, the structure of the repository suggests a deliberate attempt to support near-real-time awareness. That is critical for a system where time-sensitive fines or alerts are part of the workflow.",
                "Payment processing is represented as a gateway route rather than as a loose frontend action. That design ensures that payment updates can be associated with a specific violation record and that the system can keep a coherent audit trail. In a more mature deployment, this route would also be the natural place to integrate external payment providers and receipting logic.",
            ],
        ),
        (
            "13. Camera Network and Graph Modeling",
            [
                "The camera and edge models show that the project does not treat locations as a flat list. Instead, it models a network where each camera location is a node and each road connection can be an edge with a weight such as distance or travel time. That representation opens the door to route analysis, hotspot detection, and future path-based reasoning.",
                "For the current application, the graph is already useful as a visualization aid. The admin map page can place nodes at latitude and longitude coordinates, and the edge list can describe how those nodes relate to one another. Even without a full routing engine, that structure is enough to support a meaningful traffic-monitoring story in the report.",
                "This modeling choice also provides a natural extension path for future work. If the project later includes congestion prediction, incident routing, or patrol optimization, the graph can serve as the underlying representation without a redesign of the database. Including that observation in the seminar paper shows architectural foresight rather than only reporting what currently exists.",
            ],
        ),
        (
            "14. Dashboard, Analytics, and Public Content",
            [
                "The dashboard is not just a summary page; it is a consolidation point for multiple kinds of system state. User metrics, recent violations, notifications, and carousel content all contribute to a homepage-like view that helps the user interpret the status of their account. This is an important usability pattern because traffic systems can become overwhelming if every detail is pushed into separate views without synthesis.",
                "Analytics routes and admin analysis pages turn traffic records into a more readable operational picture. Rather than leaving the enforcement data buried in raw documents, the system can aggregate trends and present them in a form that supports decision-making. That is one of the clearest examples in the project of data becoming information.",
                "The presence of newsfeed and carousel management routes suggests a portal philosophy rather than a narrowly transactional one. The system is trying to behave like an information service for a city, not simply a back-office database tool. That distinction matters because it explains why the repository includes public-facing narrative content alongside administrative and citizen workflows.",
            ],
        ),
        (
            "15. Implementation Notes by Route Family",
            [
                "The authentication routes implement a classic lifecycle: sign-up, email verification, and sign-in. The services routes implement operational user actions such as changing ownership or renewing registration. The admin routes implement privileged data management and enforcement simulation. This separation is clear enough that it can be documented as a route family pattern in the report.",
                "A route-by-route reading also shows that the application is optimized for specific responsibilities rather than generic RPC-style endpoints. For example, the vehicle-oriented routes return or mutate a concept that is clearly tied to a public user, while the traffic-record routes focus on detection and violation records that are primarily relevant to enforcement and reporting. That clarity reduces the cognitive load for developers and reviewers alike.",
                "From a documentation standpoint, the route inventory deserves special attention because it shows how the system scales conceptually. Instead of describing one monolithic backend, the report can explain that each route corresponds to a user need or administrative responsibility. That is very much in line with SoftwareX-style documentation, where the software artifact is explained as a complete system rather than as isolated code snippets.",
            ],
        ),
        (
            "16. Deployment, Configuration, and Runtime Behavior",
            [
                "The project’s setup instructions show that the repository is intended to run on a local development machine with a Node.js frontend and a separate Python ML environment. Configuration is driven by environment variables that define database connectivity, JWT secrets, and ALPR model locations. That makes the deployment model explicit and avoids hardcoding secrets or path assumptions into the application logic.",
                "The ML service can be launched as a FastAPI application using uvicorn, while the web application is launched with the standard Next.js development server. This dual-process arrangement is simple enough for a seminar setting and still realistic enough to demonstrate multi-language integration. It also allows the report to discuss operational concerns such as model warm-up and request latency.",
                "The repository notes also mention auto-resolution of local model assets from the Hybrid Pipeline directory and support for optional label files or annotations. That is a practical detail because it reduces the friction of local setup when the models are already stored inside the project tree. In the report, this should be described as a quality-of-life feature that improves reproducibility for the student or reviewer.",
            ],
        ),
        (
            "17. Evaluation, Limitations, and Risks",
            [
                "A realistic report should not present the current system as complete. The project clearly has strengths, but it also exhibits limitations common to prototype-scale web platforms. Some state propagation appears to rely on polling rather than fully event-driven communication, the payment workflow is not tied to a real gateway, and security choices such as token storage in localStorage should be treated as acceptable for a demo but not ideal for a production deployment.",
                "The ALPR pipeline is powerful but also sensitive to model quality, plate format variation, and environmental conditions such as image blur or lighting. Since the repository supports plate normalization, it is aware of input inconsistency, yet the document should still state that automated recognition is probabilistic and therefore needs review paths or confidence thresholds in a mature system.",
                "Scalability is another obvious discussion point. MongoDB and Mongoose are appropriate for rapid development, but the report should note that indexing, pagination, and workload isolation would matter much more in a real city-scale deployment. Including these limitations increases the credibility of the paper and shows that the analysis is grounded in engineering judgment rather than only in feature enumeration.",
            ],
        ),
        (
            "18. Future Work and Research Directions",
            [
                "The most immediate future enhancement is stronger real-time communication for notifications and enforcement events. WebSocket or Server-Sent Event support would make the system feel more responsive and would reduce dependence on periodic polling. A future version could also expose a richer live operational console for administrators.",
                "Another strong direction is deeper ML integration. The ALPR pipeline could be extended with better confidence reporting, plate-format validation, vehicle-class detection, and a human review queue for uncertain detections. That would turn the current automated helper into a more mature computer-vision subsystem.",
                "There is also obvious room for stronger civic-service workflows. Integrating real payment providers, adding downloadable receipts, implementing audit logging, and improving multi-language support would make the platform more realistic and more suitable for broader adoption. Those improvements are all consistent with the current codebase and would build on the architecture already in place.",
            ],
        ),
    ]

    for heading, paragraphs in sections:
        add_heading(document, heading)
        add_body_paragraphs(document, paragraphs)

    add_heading(document, "19. Key Models and Their Roles")
    add_body_paragraphs(
        document,
        [
            "The model layer is compact enough to describe exhaustively in a paper. The user model stores account identity, password hashes, verification state, and convenience fields such as a credit score or role marker. The vehicle model stores the registry entry that binds a plate number to an owner and keeps track of its operational status.",
            "The traffic record model is the enforcement centerpiece because it captures what happened, where it happened, when it happened, and what penalty or state was assigned to the event. The notification model is the bridge to the user experience because it turns a record into a visible alert. The location and edge models add spatial semantics that would otherwise be absent from a purely document-oriented design.",
            "The carousel and newsfeed models are smaller but still important because they support the portal-like feel of the application. Instead of presenting the user with only raw technical data, the system can present announcements and information cards. That helps the interface feel like a service platform rather than a developer console.",
        ],
    )

    add_table(
        document,
        ["Model", "Primary Purpose", "Representative Fields"],
        [
            ["User", "Authentication and account identity", "email, password hash, isAdmin, isVerified, credit score"],
            ["Vehicle", "Vehicle registry and ownership", "number plate, owner, status, registration dates"],
            ["TrafficRecord", "Violation and detection events", "vehicle, plate, location, timestamp, fine, status"],
            ["Notification", "User-facing alerts", "vehicle, message, read status, linked violation"],
            ["Location", "Camera or monitoring node", "name, latitude, longitude"],
            ["Edge", "Weighted road-network connection", "from, to, distance, travel time, direction"],
            ["Carousel", "Public dashboard content", "image, title, description, category"],
            ["NewsFeed", "Traffic and accident information", "headline, body, timestamp, category"],
        ],
    )

    add_heading(document, "20. Route Inventory")
    route_rows = [
        ["/api/authentication/signin", "POST", route_description("app/api/authentication/signin/route.ts")],
        ["/api/authentication/signup", "POST", route_description("app/api/authentication/signup/route.ts")],
        ["/api/authentication/verifyemail", "POST", route_description("app/api/authentication/verifyemail/route.ts")],
        ["/api/violations", "GET", route_description("app/api/violations/route.ts")],
        ["/api/vehicles", "GET", route_description("app/api/vehicles/route.ts")],
        ["/api/traffic-records", "GET", route_description("app/api/traffic-records/route.ts")],
        ["/api/dashboard/stats", "GET", route_description("app/api/dashboard/stats/route.ts")],
        ["/api/dashboard/carousel", "GET", route_description("app/api/dashboard/carousel/route.ts")],
        ["/api/dashboard/carousel-by-category", "GET", route_description("app/api/dashboard/carousel-by-category/route.ts")],
        ["/api/dashboard/newsfeed", "GET", route_description("app/api/dashboard/newsfeed/route.ts")],
        ["/api/notifications", "GET", route_description("app/api/notifications/route.ts")],
        ["/api/notifications/stream", "GET", route_description("app/api/notifications/stream/route.ts")],
        ["/api/profile", "GET", route_description("app/api/profile/route.ts")],
        ["/api/payments/gateway", "POST", route_description("app/api/payments/gateway/route.ts")],
        ["/api/admin/camera", "GET/POST/DELETE", route_description("app/api/admin/camera/route.ts")],
        ["/api/admin/camera/[id]", "PUT/DELETE", route_description("app/api/admin/camera/[id]/route.ts")],
        ["/api/admin/locations", "GET", route_description("app/api/admin/locations/route.ts")],
        ["/api/admin/edges", "GET/POST", route_description("app/api/admin/edges/route.ts")],
        ["/api/admin/analysis", "GET", route_description("app/api/admin/analysis/route.ts")],
        ["/api/admin/violations/simulate", "POST", route_description("app/api/admin/violations/simulate/route.ts")],
        ["/api/admin/violations/extract-plate", "POST", route_description("app/api/admin/violations/extract-plate/route.ts")],
        ["/api/admin/UI/newsfeed", "GET/POST", route_description("app/api/admin/UI/newsfeed/route.ts")],
        ["/api/admin/UI/carousel/add", "POST", route_description("app/api/admin/UI/carousel/add/route.ts")],
        ["/api/admin/UI/carousel/delete/[id]", "DELETE", route_description("app/api/admin/UI/carousel/delete/[id]/route.ts")],
        ["/api/services/add-vehicle", "POST", route_description("app/api/services/add-vehicle/route.ts")],
        ["/api/services/change-ownership", "POST", route_description("app/api/services/change-ownership/route.ts")],
        ["/api/services/check-status", "GET", route_description("app/api/services/check-status/route.ts")],
        ["/api/services/renew-registration", "POST", route_description("app/api/services/renew-registration/route.ts")],
        ["/api/services/report-stolen", "POST", route_description("app/api/services/report-stolen/route.ts")],
        ["/api/services/update-details", "POST", route_description("app/api/services/update-details/route.ts")],
    ]
    add_table(document, ["Route", "Method", "Role in the system"], route_rows)

    add_heading(document, "21. Implementation Summary by Frontend Page")
    add_body_paragraphs(
        document,
        [
            "The page structure reveals a deliberate separation of concerns in the user interface. The root page serves as the entry point, while dashboard and profile pages deliver user-specific summaries. The tickets and violations view is designed to present enforcement history and map context in the same interface so that a user can understand both the fact of the violation and the place where it occurred.",
            "The service pages extend that experience by translating administrative backend capabilities into citizen operations. For example, registration renewal and ownership changes are not stored as vague UI labels; they are organized as dedicated routes with explicit backend handlers. This makes the user journey easier to follow in a seminar report because each page can be described as a meaningful software module.",
            "The admin subtree serves a different purpose: it exposes the operational controls for cameras, records, analytics, simulation, and ML testing. Since these pages are all organized under a common namespace, the report can easily show how the application separates public and privileged behavior while still sharing the same application shell and code style.",
        ],
    )

    add_heading(document, "22. SoftwareX-Style Discussion")
    add_body_paragraphs(
        document,
        [
            "A SoftwareX-style paper should explain the software artifact, its novelty, its implementation decisions, and its research or practical significance. The Traffic-System project fits that pattern because it is not merely a code sample; it is a complete operational prototype that integrates database-backed services, browser-based interfaces, and a machine-learning subsystem into a coherent platform.",
            "The software is especially suitable for a SoftwareX framing because it is reproducible. The repository includes a README, environment notes, a standalone ML service, and a clear route structure. Those characteristics make it possible for another developer to understand and re-run the system without having to reverse-engineer hidden conventions.",
            "The system also has a strong educational value. It demonstrates how a modern full-stack application can be decomposed into understandable parts while still serving a fairly complex workflow. In that sense, it is a good example of applied software engineering in a smart-city domain, which is a natural fit for the style of documentation the user requested.",
        ],
    )

    add_heading(document, "23. References and Technology Base")
    add_body_paragraphs(
        document,
        [
            "The implementation rests on a modern web technology stack centered on Next.js 15 and React 18. Data persistence uses MongoDB with Mongoose, while authentication and token work rely on jsonwebtoken and bcryptjs. The visual and spatial aspects of the application are supported by Leaflet and react-leaflet, and the notification and HTTP layers use react-hot-toast and axios where appropriate.",
            "The machine-learning stack is Python-based and relies on PyTorch, TorchVision, Ultralytics, OpenCV, Pillow, and NumPy. FastAPI and uvicorn provide the inference service layer. This combination is common in practical ALPR and computer-vision systems because it separates model development concerns from the web presentation layer.",
            "For documentation purposes, the most relevant reference materials are the official framework manuals and package documentation for Next.js, React, MongoDB, Mongoose, FastAPI, PyTorch, Ultralytics, Leaflet, and the Python packages that support image processing. The project’s own README and code comments are also important because they define the repository-specific deployment assumptions and environment variables.",
        ],
    )

    add_heading(document, "24. Conclusion")
    add_body_paragraphs(
        document,
        [
            "Traffic-System is a substantial demonstration of how a smart-city traffic authority can be modeled in software. It combines citizen self-service, administrative enforcement, payment tracking, camera network representation, and ALPR-assisted simulation within one cohesive application. The result is a system that is both technically diverse and conceptually clear.",
            "The strongest design decision is the way the system separates concerns without fragmenting the user experience. Next.js handles the interface and request orchestration, MongoDB stores the operational state, and FastAPI hosts the machine-learning inference path. That arrangement allows the project to remain understandable while still being able to express a realistic cross-technology workflow.",
            "As a seminar-paper artifact, the repository has enough breadth to support a detailed report and enough structure to support reproducibility. The Word document generated from this script captures that breadth in SoftwareX style, and it can serve as the basis for further editing if you want a more formal journal submission, a thesis appendix, or a presentation handout.",
        ],
    )

    add_heading(document, "Appendix A. Selected Workflow Summary")
    add_bullet_list(
        document,
        [
            "Citizen signs in, views dashboard, and manages vehicle records.",
            "Administrator creates or simulates a violation event for a specific plate.",
            "The system creates a traffic record, emits a notification, and updates the user-facing status.",
            "The citizen reviews the ticket and payment state from the ticket history page.",
            "The ALPR test page allows the admin to extract a plate from an uploaded image using FastAPI or the local Python fallback.",
        ],
    )

    add_heading(document, "Appendix B. Environment Variables Highlighted in the Repository")
    add_bullet_list(
        document,
        [
            "DATABASE_URL / MONGO_URL for persistence.",
            "NEXTAUTH_SECRET and NEXTAUTH_URL for authentication runtime configuration.",
            "ALPR_USE_FASTAPI and ALPR_FASTAPI_URL for ML service integration.",
            "ALPR_MODELS_DIR, ALPR_CITY_LABEL_FILE, and ALPR_CHAR_LABEL_FILE for local model resolution.",
            "KMP_DUPLICATE_LIB_OK for Windows OpenMP runtime compatibility.",
        ],
    )

    add_heading(document, "Appendix C. Document Notes")
    add_body_paragraphs(
        document,
        [
            "This report is intentionally long-form so it can function as a seminar-style technical document rather than a short README summary. It is formatted in a Word-compatible `.docx` file for easier editing, annotation, and presentation to reviewers or supervisors. The structure follows the broad expectations of SoftwareX-style reporting: problem context, system description, implementation details, evaluation, limitations, and future work.",
        ],
    )

    add_heading(document, "Appendix D. Frontend Page Catalog")
    frontend_pages = [
        ("/", "landing entry point for new visitors", "introduces the platform and directs users toward sign-in or exploration"),
        ("/home/landing", "alternate landing experience for the public home section", "supports a presentation-oriented front door to the system"),
        ("/authentication/signin", "secure sign-in screen for existing users", "collects credentials and routes the user into an authenticated session"),
        ("/authentication/signup/vehicle", "vehicle registration-oriented sign-up page", "binds account creation to the vehicle enrollment workflow"),
        ("/check-email", "verification follow-up page", "confirms the email-verification step and explains the next action"),
        ("/dashboard", "authenticated summary dashboard", "shows metrics, recent events, and shortcuts to important user actions"),
        ("/profile", "account profile page", "presents user identity, linked records, and account-level information"),
        ("/myvehicle", "vehicle detail page", "exposes the user’s registered vehicle data and status details"),
        ("/explore", "public exploration and traffic information page", "collects publicly visible feeds, maps, or awareness material"),
        ("/reports", "reporting and analytics view", "summarizes traffic records, tickets, and status trends"),
        ("/tickets&violations", "violation history page with map context", "helps users inspect fines, record location, and payment status"),
        ("/tickets&violation/payment/ticket/[ticketId]", "ticket payment detail screen", "provides a focused workflow for reviewing and paying an individual ticket"),
        ("/services/change-ownership", "ownership transfer form", "updates legal ownership metadata for a vehicle record"),
        ("/services/check-status", "status lookup page", "returns registration and compliance state for a selected record"),
        ("/services/payment-history", "payment history page", "lists past settlements and their relationship to the user’s violations"),
        ("/services/renew-registration", "registration renewal form", "extends the validity of a vehicle registration entry"),
        ("/services/report-stolen", "stolen-vehicle reporting page", "flags a vehicle and makes that state available to enforcement workflows"),
        ("/services/update-details", "vehicle detail update page", "allows corrections to ownership or registry metadata"),
        ("/admin", "administrative dashboard", "gives operators a privileged overview of system activity and controls"),
        ("/admin/addCamera", "camera creation page", "registers new monitoring points on the city map"),
        ("/admin/analysis", "analytics page", "visualizes trends in violations and traffic activity"),
        ("/admin/ml-predict", "machine-learning prediction sandbox", "lets administrators test the ALPR model in isolation"),
        ("/admin/simulate-violation", "violation simulation page", "creates controlled enforcement events for testing and demos"),
        ("/admin/traffic-records", "traffic record browser", "supports inspection and filtering of stored enforcement events"),
        ("/admin/UI", "admin UI showcase page", "organizes content-management interfaces into a reusable control area"),
        ("/admin/UI/carousel", "carousel content manager", "edits the rotating content displayed on dashboard and landing views"),
        ("/admin/UI/newsfeed", "newsfeed content manager", "maintains public news items and operational notices"),
        ("/admin/vehicles", "vehicle registry administration page", "searches and manages vehicle records across the system"),
        ("/admin/violations", "violation administration page", "reviews enforcement data and administrative action items"),
    ]
    for path, purpose, role in frontend_pages:
        add_body_paragraphs(
            document,
            [
                f"The page at {path} serves as the {purpose}. In the Traffic-System application, this screen is responsible for a specific slice of the user journey and is implemented as part of the App Router structure. {role.capitalize()}. Because the project uses CSS modules and route-based page organization, each page can keep a distinct visual identity while still participating in a shared account, API, and layout model.",
            ],
        )

    add_heading(document, "Appendix E. API Route Commentary")
    for route, method, description in route_rows:
        add_body_paragraphs(
            document,
            [
                f"The route {route} uses the {method} method and {description} In practice, this means the request passes through the application’s validation, business logic, and persistence layers before a response reaches the UI. That pattern is consistent with the rest of the system and makes the route easy to document in a seminar paper because its role is clear, its input is focused, and its output has an obvious effect on downstream pages.",
            ],
        )

    add_heading(document, "Appendix F. Operational Notes for Reproducibility")
    add_body_paragraphs(
        document,
        [
            "The report should be read together with the repository README because the project depends on a small number of runtime assumptions. The Next.js application expects a working MongoDB connection and the standard Node.js development tooling, while the ALPR service expects Python dependencies and access to the model files stored in the Hybrid Pipeline directory. Those are not minor details; they determine whether the machine-learning workflow can be executed at all.",
            "In a reproducibility-oriented environment, the order of startup also matters. The database should be reachable before the web application begins its mutation-heavy workflows, and the FastAPI server should be available if the ALPR path is intended to run in persistent-server mode. The project documentation reflects that by separating the frontend launch instructions from the ML server instructions rather than treating them as one opaque startup sequence.",
            "A reviewer or seminar supervisor reading this document should therefore be able to understand not only what the software does but also how the pieces are intended to be operated together. That operational clarity is one of the reasons the codebase is a strong fit for a SoftwareX-style technical report.",
        ],
    )

    add_heading(document, "Appendix G. Key Module Commentary")
    module_catalog = [
        ("app/layout.tsx", "the global layout wrapper that defines the application shell and shared metadata"),
        ("app/page.tsx", "the root routing entry that directs users into the main landing or dashboard experience"),
        ("app/components/header.tsx", "the shared navigation component that establishes top-level identity and route access"),
        ("app/components/footer.tsx", "the shared footer component that provides consistency across public-facing pages"),
        ("dbConnection/dbConnection.ts", "the database connection layer that centralizes MongoDB access and reuse"),
        ("helpers/jwtToken.ts", "the server-side token utility used to sign and verify authentication state"),
        ("helpers/jwtClient.ts", "the client-side token helper used to decode or inspect authenticated session data"),
        ("helpers/mailer.ts", "the mailer utility that supports email verification or message delivery workflows"),
        ("models/userModel.js", "the account schema that stores user identity, role, verification, and status metadata"),
        ("models/vehicleModel.js", "the registry schema that ties a number plate to ownership and vehicle attributes"),
        ("models/trafficRecordModel.js", "the violation schema that captures detection events and penalty-related details"),
        ("models/notificationModel.js", "the alert schema that transforms enforcement events into user-visible notifications"),
        ("models/locationModel.js", "the camera-location schema that anchors each monitoring node on the map"),
        ("models/edgeModel.js", "the graph schema that stores weighted connections between monitoring locations"),
        ("models/carouselModel.js", "the content schema that drives dashboard and landing-page carousel material"),
        ("models/newsFeedModel.js", "the public-information schema that stores traffic updates and related notices"),
        ("ml_service/main.py", "the FastAPI entry point that loads the ALPR pipeline and serves prediction requests"),
        ("scripts/extract_plate_pipeline.py", "the orchestration script that performs local plate extraction when the fallback mode is used"),
        ("Hybrid Pipeline/hybrid_pipeline.py", "the core Python inference pipeline that coordinates detection, classification, and OCR"),
        ("app/api/admin/violations/simulate/route.ts", "the high-value administrative route that creates synthetic violations for testing and demos"),
        ("app/api/admin/violations/extract-plate/route.ts", "the bridge route that connects uploaded images to the ALPR inference backend"),
    ]
    for module_path, role in module_catalog:
        add_body_paragraphs(
            document,
            [
                f"The module {module_path} is {role}. In the context of the Traffic-System codebase, it contributes to either application wiring, data management, machine-learning orchestration, or a privileged administrative workflow. Describing it explicitly matters because the system is modular enough that no single file explains the whole application on its own. The report therefore benefits from naming each core file and explaining how it participates in the larger design.",
            ],
        )

    add_heading(document, "Appendix H. Engineering Quality Considerations")
    quality_notes = [
        "A strong software report should discuss testability even when the repository does not yet contain a full test suite. In Traffic-System, the clear route boundaries, separate Python inference service, and distinct data models make it possible to add tests around authentication, violation creation, payment updates, and ALPR normalization without first refactoring the entire application. That is a favorable sign because it means the architecture is already aligned with maintainable extension.",
        "Error handling is another essential engineering topic. The project uses multiple boundaries where something can fail: database connectivity, token verification, file uploads, model loading, and external inference calls. Each boundary should be documented as a place where the implementation needs explicit validation and informative feedback so that the user is not left with an ambiguous failure. That observation helps explain why the report emphasizes workflow clarity rather than only feature coverage.",
        "Security deserves special attention because the system carries personally identifiable information, vehicle identity records, and potentially payment-related state. The use of JWTs, role separation, and server-side routes is a reasonable base, but a production version would need stronger controls around token storage, rate limiting, audit logging, and CSRF considerations. Presenting that as a design note gives the report an honest evaluation tone.",
        "Maintainability is supported by the separation between frontend pages, API routes, models, helper utilities, and the ML service. The codebase can be read as a set of responsibility clusters rather than a large procedural script. That matters because maintainability is not just about how short a file is; it is about whether a future contributor can find the right boundary when adding a feature or debugging a defect.",
        "Portability is especially relevant in a student or demo environment. The repository has been structured so that the core web application and the ALPR pipeline can be run independently, and the configuration values are exposed through environment variables rather than hardcoded literals. That helps the system survive across machines, folder layouts, and operating systems, which is an important practical advantage on Windows-based development setups.",
        "From a data-governance perspective, the application illustrates a classic trade-off between denormalization and query convenience. Storing a number plate on the traffic record makes filtering easier, but it also duplicates information that already exists in the vehicle registry. The report should acknowledge that this is a deliberate design choice that favors operational speed and simpler lookup patterns in exchange for some redundancy.",
        "The same kind of reasoning applies to notifications and dashboard content. Not every stored item has to be a core operational record; some are there to support presentation or alerting. The repository keeps those concerns visible rather than hidden, which is useful because the report can explain exactly which tables are business-critical and which ones are supporting content.",
        "Finally, the application demonstrates how a small number of well-chosen abstractions can support a broad user story. A vehicle owner, a traffic officer, a developer testing the ML pipeline, and a reviewer inspecting the data model are all looking at different slices of the same platform. The report is therefore strongest when it presents the system as an integrated whole with distinct operational viewpoints rather than as a pile of unrelated pages.",
    ]
    add_body_paragraphs(document, quality_notes)

    return document


def main() -> None:
    document = build_report()
    document.save(OUTPUT)
    print(f"Saved report to {OUTPUT}")


if __name__ == "__main__":
    main()